import os
import io
import sys
import asyncio
import logging
import shutil
import tempfile
import uuid
from flask import Flask, request, send_file, render_template, jsonify

import edge_tts
import pytesseract
from PIL import Image
import cv2
import numpy as np
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Logging ──────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(name)s] %(message)s")
logger = logging.getLogger("TTS-Server")

# ── Tesseract path detection ─────────────────────────────
_tess_paths = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
]
for _p in _tess_paths:
    if os.path.isfile(_p):
        pytesseract.pytesseract.tesseract_cmd = _p
        logger.info(f"Tesseract found at {_p}")
        break
else:
    if shutil.which("tesseract"):
        logger.info("Tesseract found on PATH")
    else:
        logger.warning("Tesseract NOT found. OCR features will not work.")

# ── Flask app ────────────────────────────────────────────
app = Flask(__name__, template_folder="templates")
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB per upload

VOICES = {
    "Jenny (US - Female)":       "en-US-JennyNeural",
    "Guy (US - Male)":           "en-US-GuyNeural",
    "Aria (US - Female)":        "en-US-AriaNeural",
    "Davis (US - Male)":         "en-US-DavisNeural",
    "Jane (US - Female)":        "en-US-JaneNeural",
    "Jason (US - Male)":         "en-US-JasonNeural",
    "Sara (US - Female)":        "en-US-SaraNeural",
    "Tony (US - Male)":          "en-US-TonyNeural",
    "Amber (US - Female)":       "en-US-AmberNeural",
    "Brandon (US - Male)":       "en-US-BrandonNeural",
    "Cora (US - Female)":        "en-US-CoraNeural",
    "Christopher (US - Male)":   "en-US-ChristopherNeural",
    "Sonia (UK - Female)":       "en-GB-SoniaNeural",
    "Ryan (UK - Male)":          "en-GB-RyanNeural",
    "Libby (UK - Female)":       "en-GB-LibbyNeural",
    "Thomas (UK - Male)":        "en-GB-ThomasNeural",
    "Neerja (India - Female)":   "en-IN-NeerjaNeural",
    "Prabhat (India - Male)":    "en-IN-PrabhatNeural",
}

VOICE = "en-US-JennyNeural"
RATE = "+0%"
PITCH = "+0Hz"

# ── In-memory document state ────────────────────────────
# Each element: {"type": "heading"|"paragraph", "level": 1-6 (headings only), "text": "..."}
_active_doc = []       # list of elements
_last_audio = None
_last_text = ""
_seen_headers = set()  # Track recent headers to avoid repetition

ALLOWED_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}
MAX_IMAGES_PER_REQUEST = 10


# ═══════════════════════════════════════════════════════════
# OCR PIPELINE
# ═══════════════════════════════════════════════════════════

# Common OCR misreads to fix in post-processing
_OCR_FIXES = [
    ("pn", "on"),    # p misread as o is rare; but "pn" as standalone is almost always "on"
    ("rn", "m"),     # classic OCR: rn → m
    ("vv", "w"),     # vv → w
    ("cl", "d"),     # cl → d
    ("tbe", "the"),
    ("bave", "have"),
    ("witb", "with"),
]


def _fix_ocr_word(word):
    """Fix common single-word OCR misreads."""
    low = word.lower()
    # Only fix standalone short misreads — don't corrupt longer words
    if low == "pn":
        return "on"
    if low == "Pn":
        return "On"
    if low == "tbe":
        return "the"
    if low == "Tbe":
        return "The"
    if low == "bave":
        return "have"
    if low == "witb":
        return "with"
    return word


def _postprocess_text(text):
    """Apply word-level OCR corrections to extracted text."""
    words = text.split(" ")
    fixed = [_fix_ocr_word(w) for w in words]
    return " ".join(fixed)


def preprocess_image(pil_img):
    """Convert PIL image to preprocessed OpenCV image for OCR.
    Uses gentle preprocessing to avoid garbling characters."""
    img = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

    # Use CLAHE for contrast enhancement (gentler than thresholding)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)

    # Otsu's threshold — auto-picks optimal level, much less aggressive
    _, thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Upscale small images for better OCR accuracy
    h, w = thresh.shape
    if max(h, w) < 1500:
        scale = 1500 / max(h, w)
        thresh = cv2.resize(thresh, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    return thresh


def _get_bbox(element):
    """Extract bounding box (x1, y1, x2, y2) from an hOCR element's title attribute."""
    title = element.get("title", "")
    for part in title.split(";"):
        part = part.strip()
        if part.startswith("bbox"):
            coords = part.replace("bbox", "").strip().split()
            if len(coords) == 4:
                return [int(c) for c in coords]
    return None


def parse_hocr(hocr_html):
    """Parse hOCR output and return structured elements.

    Key logic:
    - Lines within the same ocr_par that are close together vertically
      are joined with a SPACE (they're wrapped text, not new paragraphs).
    - Only lines with a large vertical gap between them are treated as
      separate paragraphs.
    - Heading detection uses word height vs median height.
    """
    soup = BeautifulSoup(hocr_html, "lxml")
    elements = []

    # ── Collect all word heights for heading detection ────
    all_word_heights = []
    for w in soup.find_all("span", class_="ocrx_word"):
        bbox = _get_bbox(w)
        if bbox:
            h = bbox[3] - bbox[1]
            if h > 0:
                all_word_heights.append(h)

    median_h = float(np.median(all_word_heights)) if all_word_heights else 20.0
    heading_threshold = median_h * 1.4

    # ── Process all lines in the document ───────────────
    all_lines = soup.find_all(["span", "p", "div"], class_="ocr_line")
    line_data = []
    
    for line in all_lines:
        text_parts = []
        word_heights = []
        for w in line.find_all(["span", "p", "div"], class_="ocrx_word"):
            txt = w.get_text(strip=True)
            if txt:
                text_parts.append(txt)
                bbox = _get_bbox(w)
                if bbox:
                    word_heights.append(bbox[3] - bbox[1])
                    
        line_text = " ".join(text_parts).strip()
        if not line_text:
            continue
            
        line_bbox = _get_bbox(line)
        if line_bbox:
            line_data.append({
                "text": line_text,
                "bbox": line_bbox,
                "word_heights": word_heights,
            })
            
    if not line_data:
        return []

    # Sort vertically by top coordinate
    line_data.sort(key=lambda x: x["bbox"][1])
    
    # Calculate document-wide statistics
    lefts = [ld["bbox"][0] for ld in line_data]
    rights = [ld["bbox"][2] for ld in line_data]
    heights = [ld["bbox"][3] - ld["bbox"][1] for ld in line_data]
    
    median_left = float(np.median(lefts))
    median_right = float(np.median(rights))
    median_height = float(np.median(heights))
    
    # Heuristic thresholds
    indent_threshold = median_left + (median_height * 0.8)
    short_line_threshold = median_right - (median_height * 2.5)
    gap_threshold = median_height * 0.5
    
    groups = []
    current_group = [line_data[0]]
    
    for i in range(1, len(line_data)):
        prev = line_data[i - 1]
        curr = line_data[i]
        
        gap = curr["bbox"][1] - prev["bbox"][3]
        
        is_new_par = False
        
        # 1. Large vertical gap
        if gap > gap_threshold:
            is_new_par = True
        # 2. Previous line ended early (short line)
        elif prev["bbox"][2] < short_line_threshold:
            # Exception: if it ends with a comma, hyphen, or conjunction, it's likely a wrapped line, not a paragraph break
            if not prev["text"].strip().endswith((',', '-', '—')):
                is_new_par = True
        # 3. Indented first line relative to document AND previous line
        elif curr["bbox"][0] > indent_threshold and curr["bbox"][0] > prev["bbox"][0] + (median_height * 0.5):
            is_new_par = True
            
        if is_new_par:
            groups.append(current_group)
            current_group = [curr]
        else:
            current_group.append(curr)
            
    groups.append(current_group)

    # ── Convert groups to elements ───────────────────
    for group in groups:
        full_text = " ".join(ld["text"] for ld in group)
        full_text = _postprocess_text(full_text)

        if not full_text.strip():
            continue

        all_wh = []
        for ld in group:
            all_wh.extend(ld["word_heights"])

        avg_h = float(np.mean(all_wh)) if all_wh else 0
        is_heading = avg_h > heading_threshold and len(full_text.strip()) < 200

        if is_heading:
            ratio = avg_h / median_h
            if ratio > 2.2:
                level = 1
            elif ratio > 1.8:
                level = 2
            else:
                level = 3
            elements.append({"type": "heading", "level": level, "text": full_text.strip()})
        else:
            elements.append({"type": "paragraph", "text": full_text.strip()})

    # Fallback if hOCR parsing produced nothing
    if not elements:
        raw = soup.get_text(separator="\n").strip()
        if raw:
            for block in raw.split("\n\n"):
                block = block.strip()
                if block:
                    # Join lines within a block with space
                    joined = " ".join(block.split("\n"))
                    joined = _postprocess_text(joined)
                    if joined:
                        elements.append({"type": "paragraph", "text": joined})

    return elements


def ocr_image(pil_img):
    """Run full OCR pipeline on a PIL image and return structured elements."""
    preprocessed = preprocess_image(pil_img)
    pil_preprocessed = Image.fromarray(preprocessed)

    # Get hOCR output for layout analysis
    try:
        hocr_bytes = pytesseract.image_to_pdf_or_hocr(
            pil_preprocessed, extension="hocr",
            config="--psm 3 --oem 3"
        )
        hocr_html = hocr_bytes.decode("utf-8")
        elements = parse_hocr(hocr_html)
    except Exception as e:
        logger.warning(f"hOCR parsing failed, falling back to plain text: {e}")
        # Fallback to plain text
        text = pytesseract.image_to_string(
            pil_preprocessed, config="--psm 3 --oem 3"
        )
        elements = []
        for block in text.split("\n\n"):
            block = block.strip()
            if block:
                # Join wrapped lines with space, not newline
                joined = " ".join(block.split("\n"))
                joined = _postprocess_text(joined)
                if joined:
                    elements.append({"type": "paragraph", "text": joined})

    return elements


# ═══════════════════════════════════════════════════════════
# DOCUMENT MANAGEMENT
# ═══════════════════════════════════════════════════════════

def build_docx(elements):
    """Build a .docx Document from a list of elements."""
    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for el in elements:
        if el["type"] == "heading":
            level = el.get("level", 1)
            doc.add_heading(el["text"], level=min(level, 9))
        else:
            doc.add_paragraph(el["text"])

    return doc


def parse_docx(file_stream):
    """Parse a .docx file and return structured elements."""
    doc = DocxDocument(file_stream)
    elements = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            # Extract level from style name like "Heading 1", "Heading 2"
            level = 1
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break
            elements.append({"type": "heading", "level": level, "text": text})
        else:
            elements.append({"type": "paragraph", "text": text})
    return elements


# ═══════════════════════════════════════════════════════════
# TTS ENGINE
# ═══════════════════════════════════════════════════════════

def _generate_tts(text, voice=None, rate=None):
    async def _run():
        v = voice or VOICE
        r = rate or RATE
        communicate = edge_tts.Communicate(text, v, rate=r, pitch=PITCH)
        mp3_data = io.BytesIO()
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                mp3_data.write(chunk["data"])
        mp3_data.seek(0)
        return mp3_data
    return asyncio.run(_run())


# ═══════════════════════════════════════════════════════════
# ROUTES
# ═══════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/speak", methods=["POST"])
def speak():
    global _last_audio, _last_text
    data = request.get_json()
    if not data or "text" not in data:
        return {"error": "No text provided"}, 400

    text = data["text"].strip()
    if not text:
        return {"error": "Empty text"}, 400

    voice = data.get("voice")
    rate = data.get("rate")
    voice_short = VOICES.get(voice, voice or VOICE)

    logger.info(f"Generating speech: '{text[:80]}...' ({len(text)} chars) voice={voice_short} rate={rate}")

    try:
        mp3_data = _generate_tts(text, voice=voice_short, rate=rate)
        logger.info(f"Generated {mp3_data.getbuffer().nbytes} bytes of audio")
        _last_audio = io.BytesIO(mp3_data.getvalue())
        _last_text = text
        return send_file(mp3_data, mimetype="audio/mpeg")
    except Exception as e:
        logger.error(f"TTS generation failed: {e}")
        return {"error": str(e)}, 500


@app.route("/voices")
def list_voices():
    return {"voices": list(VOICES.keys())}


@app.route("/download")
def download():
    if _last_audio is None:
        return {"error": "No audio generated yet"}, 404
    _last_audio.seek(0)
    download_name = (_last_text[:40].strip().replace(" ", "_") or "speech") + ".mp3"
    return send_file(_last_audio, mimetype="audio/mpeg", as_attachment=True, download_name=download_name)


# ── Continuity utilities ───────────────────────────────────
import re

def process_page_elements(elements):
    """Strip footers and repetitive headers from a single page's elements."""
    if not elements:
        return []
        
    # Footer: Check if last few elements are page numbers or very short noise
    # We check up to the last 2 elements just in case Tesseract picked up a smudge below the footer
    for _ in range(2):
        if not elements:
            break
        last_text = elements[-1]["text"].strip()
        if re.match(r'^(\d+|page\s*\d+|[ivxlcdm]+)$', last_text, re.IGNORECASE) or len(last_text) <= 3:
            elements.pop()
        else:
            break
            
    # Header: Check if first element is short and repetitive
    if elements:
        first_text = elements[0]["text"].strip()
        if len(first_text) < 60:
            if first_text.lower() in _seen_headers:
                elements.pop(0)
            else:
                _seen_headers.add(first_text.lower())
                
    return elements

def merge_element_lists(list1, list2):
    """Merge two lists of elements, combining boundary paragraphs if they flow together."""
    if not list1: return list2
    if not list2: return list1
    
    res = list(list1)
    incoming = list(list2)
    
    if res[-1]["type"] == "paragraph" and incoming[0]["type"] == "paragraph":
        last_par = res[-1]["text"].strip()
        first_par = incoming[0]["text"].strip()
        
        ends_with_terminator = last_par.endswith(('.', '!', '?', '"', "'", '”', '’'))
        starts_with_lowercase = first_par[0].islower() if first_par else False
        
        if not ends_with_terminator or starts_with_lowercase:
            # Merge!
            res[-1]["text"] = last_par + " " + first_par
            incoming.pop(0)
            
    res.extend(incoming)
    return res


# ── OCR endpoint ─────────────────────────────────────────
@app.route("/ocr", methods=["POST"])
def ocr_endpoint():
    if "images" not in request.files:
        return jsonify({"error": "No images uploaded"}), 400

    files = request.files.getlist("images")
    if len(files) > MAX_IMAGES_PER_REQUEST:
        return jsonify({"error": f"Maximum {MAX_IMAGES_PER_REQUEST} images allowed"}), 400

    all_elements = []
    for f in files:
        ext = os.path.splitext(f.filename or "")[1].lower()
        if ext not in ALLOWED_IMAGE_EXT:
            return jsonify({"error": f"Unsupported image type: {ext}"}), 400

        try:
            pil_img = Image.open(f.stream)
            elements = ocr_image(pil_img)
            elements = process_page_elements(elements)
            
            # Merge elements within this batch seamlessly
            all_elements = merge_element_lists(all_elements, elements)
            logger.info(f"OCR processed {f.filename}")
        except Exception as e:
            logger.error(f"OCR failed for {f.filename}: {e}")
            return jsonify({"error": f"OCR failed for {f.filename}: {str(e)}"}), 500

    return jsonify({"elements": all_elements})


# ── Document endpoints ───────────────────────────────────
@app.route("/doc/append", methods=["POST"])
def doc_append():
    global _active_doc
    data = request.get_json()
    if not data or "elements" not in data:
        return jsonify({"error": "No elements provided"}), 400
        
    _active_doc = merge_element_lists(_active_doc, data["elements"])
    logger.info(f"Appended elements. Doc now has {len(_active_doc)} elements.")
    return jsonify({"ok": True, "total": len(_active_doc)})


@app.route("/doc/content")
def doc_content():
    return jsonify({"elements": _active_doc})


@app.route("/doc/save")
def doc_save():
    if not _active_doc:
        return jsonify({"error": "Document is empty"}), 400
    doc = build_docx(_active_doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="document.docx")


@app.route("/doc/new", methods=["POST"])
def doc_new():
    global _active_doc, _seen_headers
    _active_doc = []
    _seen_headers.clear()
    logger.info("Document cleared")
    return jsonify({"ok": True})


@app.route("/doc/open", methods=["POST"])
def doc_open():
    global _active_doc
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    ext = os.path.splitext(f.filename or "")[1].lower()
    if ext != ".docx":
        return jsonify({"error": "Only .docx files are supported"}), 400
    try:
        elements = parse_docx(f.stream)
        _active_doc = elements
        logger.info(f"Opened document with {len(elements)} elements")
        return jsonify({"elements": elements})
    except Exception as e:
        logger.error(f"Failed to open document: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/doc/text")
def doc_text():
    """Get all document text concatenated for TTS."""
    if not _active_doc:
        return jsonify({"error": "Document is empty"}), 400
    parts = []
    for el in _active_doc:
        parts.append(el["text"])
    return jsonify({"text": "\n\n".join(parts)})


# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5070))
    logger.info(f"Starting TTS server on http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=False)
